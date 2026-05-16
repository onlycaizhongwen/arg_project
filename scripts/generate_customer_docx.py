from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path("docs/codex/v1/plans/data-cleaning-rag-mvp-customer-technical-proposal.md")
OUTPUT = Path("docs/codex/v1/plans/数据清洗与RAG服务MVP技术方案建议书.docx")


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name].font.size = Pt(size)


def add_table(doc: Document, lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index in range(max_cols):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[col_index].replace("`", "") if col_index < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9.5)
                    if row_index == 0:
                        run.bold = True
    doc.add_paragraph()


def add_code(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.space_after = Pt(6)


def add_image(doc: Document, markdown_line: str, base_dir: Path) -> None:
    # Markdown image: ![alt](path)
    alt_start = markdown_line.find("[")
    alt_end = markdown_line.find("]")
    path_start = markdown_line.find("(", alt_end)
    path_end = markdown_line.find(")", path_start)
    if path_start < 0 or path_end < 0:
        return
    alt = markdown_line[alt_start + 1 : alt_end] if alt_start >= 0 and alt_end > alt_start else ""
    image_path = (base_dir / markdown_line[path_start + 1 : path_end]).resolve()
    if not image_path.exists():
        doc.add_paragraph(f"[图片缺失：{image_path}]")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.8))
    if alt:
        caption = doc.add_paragraph(alt)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    lines = text.splitlines()
    in_code = False
    code_lines: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("!["):
            add_image(doc, stripped, SOURCE.parent)
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue
        if stripped.startswith("# "):
            paragraph = doc.add_paragraph(stripped[2:].strip(), style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith(">"):
            paragraph = doc.add_paragraph(stripped.lstrip(">").strip())
            paragraph.paragraph_format.left_indent = Inches(0.18)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 2 and stripped[0].isdigit() and ". " in stripped[:4]:
            doc.add_paragraph(stripped.split(". ", 1)[1].strip(), style="List Number")
        else:
            doc.add_paragraph(stripped)
        i += 1
    add_code(doc, code_lines)
    doc.core_properties.title = "数据清洗与 RAG 服务 MVP 技术方案建议书"
    doc.core_properties.subject = "MVP 对客交流接口说明"
    doc.core_properties.author = "项目组"
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_docx()
